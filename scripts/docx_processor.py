from docx import Document

def ConvertDocument(filename):
    doc = Document(filename)

    paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]

    text = "\n".join(paragraphs)

    file = filename[6:(len(filename)-5)]
    
    with open(f"../texts/{file}.txt", "w") as f:
        f.write(text)


def ReadDocument(filename):
    doc = Document(filename)

    paragraphs = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            paragraphs.append(paragraph.text)

    print(paragraphs)


def CreateTextFile(text, filename):
    with open(f"texts/{filename}", "w", encoding="utf-8") as f:
        f.write(text)

def CreateDocument(text, filename):
    print("something")