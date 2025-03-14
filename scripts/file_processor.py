import pdfplumber
import asyncio
from docx import Document

async def DOCtoText(filename):
    doc = Document(filename)
    paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
    text = "\n".join(paragraphs)
    file = filename[6:(len(filename)-5)]
    with open(f"texts/{file}_docx.txt", "w", encoding="utf-8") as f:
        f.write(text)
    print("Document converted to txt file")

async def PDFtoText(filename):
    with pdfplumber.open(filename) as pdf:
        sentences = [page.extract_text() for page in pdf.pages]
        texts = "\n".join(sentences)

        file = filename[6:(len(filename)-4)]
        with open(f"texts/{file}_pdf.txt", "w", encoding="utf-8") as f:
            f.write(texts)
        print("PDF converted to txt file")

async def CreateTextFile(text, filename):
    with open(f"texts/{filename}", "w", encoding="utf-8") as f:
        f.write(text)
    print("Created txt file")

async def CreateDocument(text, filename):
    print("something")